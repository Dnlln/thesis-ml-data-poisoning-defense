"""Применение ГОСТовского форматирования к thesis_full.docx.

ГОСТ Р 7.0.5-2008 / ГОСТ 7.32-2017:
- Шрифт: Times New Roman, 14 pt
- Межстрочный интервал: 1,5
- Отступ первой строки абзаца: 1,25 см
- Выравнивание: по ширине
- Поля: левое 3 см, правое 1.5 см, верхнее/нижнее 2 см
- Нумерация страниц: снизу по центру
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/home/user/workspace/thesis-ml-data-poisoning-defense/docs/thesis_full.docx"
DST = "/home/user/workspace/thesis-ml-data-poisoning-defense/docs/ВКР_Методика_защиты_ML_от_отравления_данных.docx"

doc = Document(SRC)

# --- 1. Поля страницы ---
for section in doc.sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# --- 2. Стили текста ---
def set_font(run, name="Times New Roman", size=14):
    run.font.name = name
    run.font.size = Pt(size)
    # для поддержки кириллицы
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def is_heading(paragraph):
    sn = (paragraph.style.name or "").lower()
    return sn.startswith("heading") or sn.startswith("title")


# Нормальный стиль — основа
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(14)
normal.font.color.rgb = RGBColor(0, 0, 0)
pf = normal.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.first_line_indent = Cm(1.25)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# --- 3. Применяем к каждому абзацу ---
for p in doc.paragraphs:
    heading = is_heading(p)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    if heading:
        # Заголовки: жирный TNR 16, по центру, без отступа
        pf.first_line_indent = Cm(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        for run in p.runs:
            set_font(run, size=16)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        # Текст: TNR 14, 1.5, отступ 1.25, по ширине
        if p.text.strip():
            pf.first_line_indent = Cm(1.25)
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            set_font(run, size=14)

# --- 4. Шрифт в таблицах ---
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    set_font(run, size=12)

# --- 5. Нумерация страниц (снизу по центру) ---
from docx.oxml.ns import nsmap

def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # очищаем
    for run in list(paragraph.runs):
        run.text = ""

    run = paragraph.add_run()
    set_font(run, size=12)

    # Простой PAGE-поле
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


for section in doc.sections:
    add_page_number(section)

doc.save(DST)
print(f"Готово: {DST}")

import os
print(f"Размер: {os.path.getsize(DST)/1024:.1f} КБ")
